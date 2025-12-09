"""
Gerador de Documentos Preenchidos (FICHA e OP)
Sistema que preenche templates .docx e .xls/.xlsx com dados da produção
"""
from docx import Document
from openpyxl import load_workbook
from datetime import datetime
import io
import os
import tempfile
import subprocess
from typing import Dict, List, Any

def fill_docx_template(template_content: bytes, data: Dict[str, Any]) -> bytes:
    """
    Preenche template .docx com dados da produção
    
    Args:
        template_content: Conteúdo do arquivo .docx em bytes
        data: Dicionário com os dados para preencher os placeholders
        
    Returns:
        bytes: Arquivo .docx preenchido
    """
    # Carregar o template do buffer
    doc = Document(io.BytesIO(template_content))
    
    # Criar mapeamento de placeholders
    # Os placeholders podem estar em vários formatos: {{LOTE}}, {LOTE}, LOTE, etc.
    placeholders = {
        'LOTE': data.get('batch_number', ''),
        'PRODUTO': data.get('product_name', ''),
        'DATA': data.get('date', ''),
        'MATERIA_PRIMA': data.get('raw_materials_text', ''),
        'LOTE_MATERIA_PRIMA': data.get('raw_materials_batches', '')
    }
    
    # Substituir em parágrafos
    for paragraph in doc.paragraphs:
        for key, value in placeholders.items():
            # Suportar vários formatos de placeholder
            for pattern in [f'{{{{{key}}}}}', f'{{{key}}}', key]:
                if pattern in paragraph.text:
                    paragraph.text = paragraph.text.replace(pattern, str(value))
    
    # Substituir em tabelas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in placeholders.items():
                        for pattern in [f'{{{{{key}}}}}', f'{{{key}}}', key]:
                            if pattern in paragraph.text:
                                paragraph.text = paragraph.text.replace(pattern, str(value))
    
    # Salvar o documento preenchido em um buffer
    output_buffer = io.BytesIO()
    doc.save(output_buffer)
    output_buffer.seek(0)
    return output_buffer.getvalue()


def generate_op_pdf(product_data, batch_data, raw_materials_data):
    """
    Gera PDF da ORDEM DE PRODUÇÃO preenchida
    """
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4,
                           topMargin=1*cm, bottomMargin=1*cm,
                           leftMargin=1.5*cm, rightMargin=1.5*cm)
    
    elements = []
    styles = getSampleStyleSheet()
    
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=16,
        textColor=colors.HexColor('#000000'),
        spaceAfter=12,
        alignment=TA_CENTER,
        fontName='Helvetica-Bold'
    )
    
    # Título
    elements.append(Paragraph("ORDEM DE PRODUÇÃO", title_style))
    elements.append(Spacer(1, 0.5*cm))
    
    # Cabeçalho
    header_data = [
        ['OP Nº:', batch_data.get('op_number', ''), '', ''],
        ['PRODUTO:', product_data.get('name', ''), 'QUANT (Lts):', batch_data.get('planned_quantity', '')],
        ['LOTE:', batch_data.get('batch_number', ''), 'DATA:', batch_data.get('date', '')],
        ['', '', 'Hora I:', '__:__', 'Hora F:', '__:__']
    ]
    
    header_table = Table(header_data, colWidths=[3*cm, 6*cm, 3*cm, 4*cm])
    header_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#DDDDDD')),
        ('BACKGROUND', (2, 0), (2, -1), colors.HexColor('#DDDDDD')),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
        ('FONTNAME', (0, 0), (-1, -1), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
    ]))
    elements.append(header_table)
    elements.append(Spacer(1, 0.7*cm))
    
    # Tabela de Matérias-Primas
    mp_data = [['CÓD', 'MATÉRIAS-PRIMAS', 'LOTE', 'VALIDADE', '%', 'QUANT. (kg)', 'CORREÇÃO', 'TOTAL', 'OBS']]
    
    if raw_materials_data:
        total_percentage = 0
        total_quantity = 0
        
        for rm in raw_materials_data:
            quantity = rm.get('quantity', 0)
            percentage = (quantity / float(batch_data.get('planned_quantity', 1))) * 100
            
            mp_data.append([
                rm.get('code', ''),
                rm.get('name', ''),
                rm.get('supplier_batch', ''),
                rm.get('expiry_date', ''),
                f"{percentage:.2f}",
                str(quantity),
                '',  # Correção
                '',  # Total
                ''   # OBS
            ])
            
            total_percentage += percentage
            total_quantity += quantity
        
        # Linha de total
        mp_data.append(['', '', '', '', f"{total_percentage:.2f}", str(total_quantity), '', '', ''])
    
    mp_table = Table(mp_data, colWidths=[1.5*cm, 4*cm, 2*cm, 2*cm, 1.5*cm, 2*cm, 1.8*cm, 1.5*cm, 1.2*cm])
    mp_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#AAAAAA')),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 7),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('ALIGN', (4, 1), (7, -1), 'CENTER'),
    ]))
    elements.append(mp_table)
    elements.append(Spacer(1, 0.8*cm))
    
    # Seção de Embalagens
    elements.append(Paragraph("<b>EMBALAGENS</b>", styles['Heading2']))
    elements.append(Spacer(1, 0.3*cm))
    
    emb_data = [
        ['EMBALAGENS', '2LTS', '5 LTS', '20 LTS', '50 LTS', 'TB', 'SOBRA', 'PERDA'],
        ['Peso Bruto', '', '', '', '', '', '', '%'],
        ['Tara', '', '', '', '', '', '', ''],
        ['Peso Líquido', '', '', '', '', '', '', ''],
        ['Quantidade', '', '', '', '', '', '', 'Litros'],
        ['Total (Litros)', '', '', '', '', '', '', '']
    ]
    
    emb_table = Table(emb_data, colWidths=[3*cm, 2*cm, 2*cm, 2*cm, 2*cm, 1.5*cm, 2*cm, 2*cm])
    emb_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#DDDDDD')),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
        ('FONTSIZE', (0, 0), (-1, -1), 8),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('ALIGN', (1, 0), (-1, -1), 'CENTER'),
    ]))
    elements.append(emb_table)
    elements.append(Spacer(1, 1*cm))
    
    # Responsáveis e Assinaturas
    resp_data = [
        ['Responsável pela pesagem', '', 'Responsável pela preparação', '', 'Responsável pelo envase', ''],
        ['Ass.', '', 'Ass.', '', 'Ass.', ''],
        ['Nome', '', 'Nome', '', 'Nome', ''],
        ['H Inicio', '', 'H Inicio', '', 'H Inicio', ''],
        ['H Térm', '', 'H Térm', '', 'H Térm', '']
    ]
    
    resp_table = Table(resp_data, colWidths=[4*cm, 1*cm, 4*cm, 1*cm, 4*cm, 1*cm])
    resp_table.setStyle(TableStyle([
        ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
        ('FONTSIZE', (0, 0), (-1, -1), 8),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
    ]))
    elements.append(resp_table)
    elements.append(Spacer(1, 0.5*cm))
    
    # Conferente/Analista
    conf_data = [
        ['Conferente / Analista', 'Ass.:', '________________', '', 'Ass.:', '________________'],
        ['', 'Nome:', '________________', '', 'Nome:', '________________']
    ]
    
    conf_table = Table(conf_data, colWidths=[4*cm, 2*cm, 3*cm, 1*cm, 2*cm, 3*cm])
    conf_table.setStyle(TableStyle([
        ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
        ('FONTSIZE', (0, 0), (-1, -1), 8),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
    ]))
    elements.append(conf_table)
    
    # Rodapé
    elements.append(Spacer(1, 0.5*cm))
    footer_style = ParagraphStyle('Footer', parent=styles['Normal'], fontSize=8, alignment=TA_CENTER)
    elements.append(Paragraph(f"Documento gerado em: {datetime.now().strftime('%d/%m/%Y %H:%M')}", footer_style))
    
    doc.build(elements)
    buffer.seek(0)
    return buffer
